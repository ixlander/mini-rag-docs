from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader



@dataclass(frozen=True)
class Section:
    title: str
    text: str


@dataclass(frozen=True)
class Document:
    doc_id: str
    source_path: str
    source_group: str  
    title: str
    sections: List[Section]
    url: Optional[str] = None


_WHITESPACE_RE = re.compile(r"[ \t]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def _normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _WHITESPACE_RE.sub(" ", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = _BLANK_LINES_RE.sub("\n\n", s)
    return s.strip()


def _make_doc_id(path: Path, root_dir: Path) -> str:
    rel = path.resolve().relative_to(root_dir.resolve())
    return str(rel).replace("\\", "/")


def _guess_group(path: Path, root_dir: Path) -> str:
    rel = path.resolve().relative_to(root_dir.resolve()).parts
    if len(rel) >= 3 and rel[0] == "data" and rel[1] == "raw":
        return rel[2]
    return "unknown"


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def parse_markdown(md: str) -> Tuple[str, List[Section]]:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    lines = md.split("\n")

    headings: List[Tuple[int, str, int]] = []
    for i, line in enumerate(lines):
        m = _MD_HEADING_RE.match(line.strip())
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            headings.append((level, title, i))

    if not headings:
        text = _normalize_text(md)
        title = "Untitled"
        return title, [Section(title="Main", text=text)] if text else [Section(title="Main", text="")]

    doc_title = next((t for lvl, t, _ in headings if lvl == 1), headings[0][1]) or "Untitled"

    sections: List[Section] = []
    for idx, (lvl, title, start_i) in enumerate(headings):
        end_i = headings[idx + 1][2] if idx + 1 < len(headings) else len(lines)
        block = "\n".join(lines[start_i + 1 : end_i])
        block = _normalize_text(block)
        sections.append(Section(title=title or "Untitled Section", text=block))

    return doc_title, sections


def extract_text_from_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return _normalize_text("\n".join(parts))


def extract_text_from_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        t = t.strip()
        if t:
            parts.append(t)
    return _normalize_text("\n\n".join(parts))


def parse_html(html: str) -> Tuple[str, List[Section]]:
    soup = BeautifulSoup(html, "lxml")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    doc_title = "Untitled"
    if soup.title and soup.title.string:
        doc_title = soup.title.string.strip() or doc_title

    sections: List[Section] = []
    current_title = "Main"
    buf: List[str] = []

    def flush():
        nonlocal buf
        text = _normalize_text("\n".join(buf))
        if text:
            sections.append(Section(title=current_title, text=text))
        buf = []

    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        name = el.name.lower()
        txt = (el.get_text(" ", strip=True) or "").strip()
        if not txt:
            continue
        if name.startswith("h"):
            flush()
            current_title = txt
        else:
            buf.append(txt)

    flush()

    if not sections:
        text = _normalize_text(soup.get_text(" ", strip=True))
        return doc_title, [Section(title="Main", text=text)] if text else [Section(title="Main", text="")]

    if not doc_title or doc_title == "Untitled":
        doc_title = sections[0].title or doc_title

    return doc_title, sections


def parse_file(path: Path, root_dir: Path) -> Document:
    ext = path.suffix.lower()

    if ext in (".md", ".markdown"):
        raw = _read_text_file(path)
        title, sections = parse_markdown(raw)

    elif ext in (".html", ".htm"):
        raw = _read_text_file(path)
        title, sections = parse_html(raw)

    elif ext == ".txt":
        raw = _read_text_file(path)
        text = _normalize_text(raw)
        title = path.stem
        sections = [Section(title="Main", text=text)]

    elif ext == ".docx":
        text = extract_text_from_docx(path)
        title = path.stem
        sections = [Section(title="Main", text=text)]

    elif ext == ".pdf":
        text = extract_text_from_pdf(path)
        title = path.stem
        sections = [Section(title="Main", text=text)]

    else:
        raise ValueError(f"Unsupported file type: {path}")

    doc_id = _make_doc_id(path, root_dir)
    group = _guess_group(path, root_dir)

    return Document(
        doc_id=doc_id,
        source_path=str(path.resolve()),
        source_group=group,
        title=title or path.stem,
        sections=sections,
        url=None,
    )

def iter_docs(raw_root: str = "data/raw") -> List[Document]:
    root = Path(raw_root).resolve()
    if not root.exists():
        raise FileNotFoundError(f"Raw data dir not found: {root}")

    docs: List[Document] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in (".md", ".markdown", ".html", ".htm", ".txt", ".docx", ".pdf"):
            continue
        docs.append(parse_file(path, root))
    return docs

